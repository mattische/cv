import os
import re
from markdown2 import markdown
from docx import Document
from fpdf import FPDF
from bs4 import BeautifulSoup
from weasyprint import HTML

# Filnamn
INDEX_FILE = "index.md"
README_FILE = "README.md"
OUTPUT_DIR = "files"
OUTPUT_PREFIX = "MattiasSchertell"

def create_output_dir():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def read_markdown_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        return file.read()

def write_file(filepath, content):
    with open(filepath, 'w', encoding='utf-8') as file:
        file.write(content)

# från md till html
def markdown_to_html(md_content):
    return markdown(md_content)

# från html till docx
def html_to_docx(html_content, output_path):
    doc = Document()
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.descendants:
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "h3":
            doc.add_heading(element.text, level=3)
        elif element.name == "strong":
            paragraph = doc.add_paragraph()  # Skapa ett nytt paragrafobjekt
            run = paragraph.add_run(element.text)
            run.bold = True  # Ställ in fetstil
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style="List Number")
        elif element.name == "p":
            doc.add_paragraph(element.text)

    doc.save(output_path)


# från html till pdf
def html_to_pdf(html_content, output_path):
    HTML(string=html_content).write_pdf(output_path)

def generate_html(md_content):
    html_content = markdown(md_content)
    html_file = os.path.join(OUTPUT_DIR, OUTPUT_PREFIX + ".html")
    write_file(html_file, html_content)

def generate_docx(md_content):
    doc = Document()
    lines = md_content.splitlines()
    for line in lines:
        if line.startswith("# "):  # Header 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  # Header 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):  # Header 3
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    docx_file = os.path.join(OUTPUT_DIR, OUTPUT_PREFIX + ".docx")
    doc.save(docx_file)

def sanitize_content(content):
    # Ta bort tecken som inte stöds av "latin-1"
    return re.sub(r'[^\x00-\xFF]', '', content)


def generate_pdf(md_content):
    sanitized_content = sanitize_content(md_content)
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    # Lägg till en Unicode-kompatibel font
    font_path = os.path.join(os.path.dirname(__file__), "fonts")
    pdf.add_font("DejaVu", "", os.path.join(font_path, "DejaVuSans.ttf"), uni=True)  # Normal
    pdf.add_font("DejaVu", "B", os.path.join(font_path, "DejaVuSans-Bold.ttf"), uni=True)  # Bold
    pdf.add_font("DejaVu", "I", os.path.join(font_path, "DejaVuSans-Oblique.ttf"), uni=True)  # Italic
    pdf.add_font("DejaVu", "BI", os.path.join(font_path, "DejaVuSans-BoldOblique.ttf"), uni=True)  # Bold-Italic

    pdf.set_font("DejaVu", size=12)

    soup = BeautifulSoup(markdown(sanitized_content), "html.parser")
    for element in soup.descendants:
        if element.name == "h1":
            pdf.set_font("DejaVu", style="B", size=16)
            pdf.cell(200, 10, txt=element.text, ln=True, align="C")
        elif element.name == "h2":
            pdf.set_font("DejaVu", style="B", size=14)
            pdf.cell(200, 10, txt=element.text, ln=True, align="L")
        elif element.name == "h3":
            pdf.set_font("DejaVu", style="B", size=12)
            pdf.cell(200, 10, txt=element.text, ln=True, align="L")
        elif element.name == "p":
            pdf.set_font("DejaVu", size=12)
            pdf.multi_cell(0, 10, txt=element.text)
    pdf_file = os.path.join(OUTPUT_DIR, OUTPUT_PREFIX + ".pdf")
    pdf.output(pdf_file)

def update_readme(md_content):
    write_file(README_FILE, md_content)

def main():
    # Skapa output-mappen om den inte finns
    create_output_dir()

    # Läs innehåll från index.md
    md_content = read_markdown_file(INDEX_FILE)


    # Konvertera Markdown till HTML
    html_content = markdown_to_html(md_content)
    html_path = os.path.join(OUTPUT_DIR, OUTPUT_PREFIX + ".html")
    write_file(html_path, html_content)

    # Generera DOCX och PDF från HTML
    docx_path = os.path.join(OUTPUT_DIR, OUTPUT_PREFIX + ".docx")
    pdf_path = os.path.join(OUTPUT_DIR, OUTPUT_PREFIX + ".pdf")
    html_to_docx(html_content, docx_path)
    html_to_pdf(html_content, pdf_path)
    
    # Generera dokument
    #generate_html(md_content)
    #generate_docx(md_content)
    #generate_pdf(md_content)

    # Uppdatera README.md
    update_readme(md_content)

if __name__ == "__main__":
    main()
